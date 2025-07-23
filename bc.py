import os
import subprocess
import time
import uuid
import threading
import logging
import magic
import json
from flask import Flask, render_template, request, send_file, jsonify, redirect, url_for
from PIL import Image
import fitz  # PyMuPDF
import pytesseract
import cv2
import numpy as np
from docx import Document
from pdf2image import convert_from_path

# Configure Flask app
app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'static/uploads'
app.config['DOWNLOAD_FOLDER'] = 'static/downloads'
app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024 * 1024  # 10GB max file size

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Progress tracking
progress_queue = queue.Queue()
progress_data = {}

# Supported video formats
SUPPORTED_VIDEO_FORMATS = {
    'mp4': 'MP4',
    'avi': 'AVI',
    'mov': 'MOV',
    'mkv': 'MKV',
    'wmv': 'WMV',
    'flv': 'FLV',
    'webm': 'WebM'
}

# Quality settings for video compression
QUALITY_LEVELS = {
    'high': {'crf': '20', 'preset': 'medium'},
    'medium': {'crf': '28', 'preset': 'fast'},
    'low': {'crf': '32', 'preset': 'ultrafast'}
}

# Ensure directories exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['DOWNLOAD_FOLDER'], exist_ok=True)

@app.route('/')
def index():
    return render_template('index.html', 
                         video_formats=SUPPORTED_VIDEO_FORMATS,
                         has_google_ads=True)  # গুগল অ্যাডস সক্রিয়

@app.route('/progress/<task_id>')
def get_progress(task_id):
    return jsonify(progress_data.get(task_id, {'status': 'not_found'}))

@app.route('/convert', methods=['POST'])
def convert():
    if 'file' not in request.files:
        return redirect(request.url)
    
    file = request.files['file']
    if file.filename == '':
        return redirect(request.url)
    
    # Get conversion type and additional parameters
    conversion_type = request.form['conversion_type']
    output_format = request.form.get('output_format', 'mp4')  # Default to MP4
    quality = request.form.get('quality', 'medium')  # Default quality
    
    # Generate unique task ID
    task_id = str(uuid.uuid4())
    
    # Start conversion in a separate thread
    thread = threading.Thread(
        target=process_conversion,
        args=(file, conversion_type, output_format, quality, task_id)
    )
    thread.start()
    
    return jsonify({
        'task_id': task_id,
        'status': 'processing',
        'message': 'Conversion started'
    })

def process_conversion(file, conversion_type, output_format, quality, task_id):
    try:
        # Initialize progress
        progress_data[task_id] = {
            'status': 'processing',
            'progress': 0,
            'message': 'Starting conversion'
        }
        
        # Generate unique filename and save
        filename = str(uuid.uuid4()) + os.path.splitext(file.filename)[1]
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        
        output_path = None
        
        # Handle different conversion types
        if conversion_type == 'compress_video':
            output_path = compress_video_large(file_path, output_format, quality, task_id)
        elif conversion_type == 'compress_picture':
            output_path = compress_picture(file_path)
        else:
            output_path = handle_standard_conversion(file_path, conversion_type)
        
        progress_data[task_id] = {
            'status': 'completed',
            'progress': 100,
            'message': 'Conversion completed',
            'download_path': output_path
        }
        
    except Exception as e:
        logger.error(f"Conversion error for task {task_id}: {str(e)}")
        progress_data[task_id] = {
            'status': 'failed',
            'progress': 0,
            'message': f'Error: {str(e)}'
        }

def handle_standard_conversion(file_path, conversion_type):
    # Handle all non-video conversions
    if conversion_type == 'jpg_to_png':
        return jpg_to_png(file_path)
    elif conversion_type == 'png_to_jpg':
        return png_to_jpg(file_path)
    elif conversion_type == 'jpg_to_pdf':
        return jpg_to_pdf(file_path)
    elif conversion_type == 'pdf_to_jpg':
        return pdf_to_jpg(file_path)
    elif conversion_type == 'word_to_jpg':
        return word_to_jpg(file_path)
    elif conversion_type == 'jpg_to_word':
        return jpg_to_word(file_path)
    elif conversion_type == 'video_to_jpg':
        return video_to_jpg(file_path)
    else:
        raise ValueError(f"Unsupported conversion type: {conversion_type}")

def get_codec_settings(output_format):
    """Return codec settings based on output format"""
    settings = {
        'video_codec': 'libx264',
        'audio_codec': 'aac',
        'format_flags': '-movflags +faststart'
    }
    
    format_lower = output_format.lower()
    
    if format_lower == 'mp4':
        settings = {'video_codec': 'libx264', 'audio_codec': 'aac', 'format_flags': '-movflags +faststart'}
    elif format_lower == 'mov':
        settings = {'video_codec': 'libx264', 'audio_codec': 'aac', 'format_flags': '-movflags +faststart'}
    elif format_lower == 'mkv':
        settings = {'video_codec': 'libx264', 'audio_codec': 'aac', 'format_flags': ''}
    elif format_lower == 'avi':
        settings = {'video_codec': 'mpeg4', 'audio_codec': 'mp3', 'format_flags': ''}
    elif format_lower == 'wmv':
        settings = {'video_codec': 'wmv2', 'audio_codec': 'wmav2', 'format_flags': ''}
    elif format_lower == 'flv':
        settings = {'video_codec': 'flv', 'audio_codec': 'mp3', 'format_flags': ''}
    elif format_lower == 'webm':
        settings = {'video_codec': 'libvpx', 'audio_codec': 'libvorbis', 'format_flags': ''}
    
    return settings

def compress_video_large(input_path, output_format, quality, task_id):
    """Compress video with progress tracking"""
    # Get video info for progress tracking
    cap = cv2.VideoCapture(input_path)
    total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
    fps = cap.get(cv2.CAP_PROP_FPS)
    duration = total_frames / fps if fps > 0 else 0
    cap.release()
    
    # Prepare output path
    output_filename = f"{uuid.uuid4()}.{output_format.lower()}"
    output_path = os.path.join(app.config['DOWNLOAD_FOLDER'], output_filename)
    
    # Get codec settings
    codec_settings = get_codec_settings(output_format)
    
    # FFmpeg command for compression
    cmd = [
        'ffmpeg',
        '-i', input_path,
        '-c:v', codec_settings['video_codec'],
        '-preset', QUALITY_LEVELS[quality]['preset'],
        '-crf', QUALITY_LEVELS[quality]['crf'],
        '-movflags', '+faststart',
        '-y',  # overwrite output
        '-progress', 'pipe:1',  # Progress output
        '-loglevel', 'error',  # Only show errors
    ]
    
    # Add audio codec if needed
    if codec_settings['audio_codec']:
        cmd.extend(['-c:a', codec_settings['audio_codec']])
    
    # Add format-specific flags
    if output_format.lower() == 'mkv':
        cmd.extend(['-map', '0'])  # Preserve all streams
    
    # Start FFmpeg process
    process = subprocess.Popen(
        cmd,
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        universal_newlines=True
    )
    
    # Track progress
    start_time = time.time()
    last_update = 0
    
    for line in process.stdout:
        if 'out_time_ms' in line:
            # Extract current time in milliseconds
            current_time = float(line.split('=')[1]) / 1000000
            progress = (current_time / duration) * 100 if duration > 0 else 0
            
            # Update progress every second to avoid flooding
            current_time_real = time.time() - start_time
            if current_time_real - last_update >= 1:
                progress_data[task_id] = {
                    'status': 'processing',
                    'progress': min(99, int(progress)),
                    'message': f'Compressing {output_format.upper()} ({int(progress)}%)'
                }
                last_update = current_time_real
    
    # Wait for process to complete
    process.wait()
    
    if process.returncode != 0:
        raise Exception(f"Video compression to {output_format.upper()} failed")
    
    return output_path

def compress_picture(input_path):
    """Compress image with resizing"""
    img = Image.open(input_path)
    output_path = os.path.join(app.config['DOWNLOAD_FOLDER'], f"{uuid.uuid4()}.jpg")
    
    # Convert to RGB if needed
    if img.mode != 'RGB':
        img = img.convert('RGB')
    
    # Resize large images while maintaining aspect ratio
    max_size = (1920, 1080)  # HD resolution
    img.thumbnail(max_size, Image.LANCZOS)
    
    # Save with compression
    img.save(output_path, 'JPEG', quality=65, optimize=True)
    return output_path

# Image conversion functions
def jpg_to_png(input_path):
    img = Image.open(input_path)
    output_path = os.path.join(app.config['DOWNLOAD_FOLDER'], f"{uuid.uuid4()}.png")
    img.save(output_path)
    return output_path

def png_to_jpg(input_path):
    img = Image.open(input_path)
    output_path = os.path.join(app.config['DOWNLOAD_FOLDER'], f"{uuid.uuid4()}.jpg")
    img.convert('RGB').save(output_path, quality=95)
    return output_path

def jpg_to_pdf(input_path):
    img = Image.open(input_path)
    output_path = os.path.join(app.config['DOWNLOAD_FOLDER'], f"{uuid.uuid4()}.pdf")
    img.convert('RGB').save(output_path, "PDF", resolution=100.0)
    return output_path

def pdf_to_jpg(input_path):
    images = convert_from_path(input_path)
    output_path = os.path.join(app.config['DOWNLOAD_FOLDER'], f"{uuid.uuid4()}.jpg")
    images[0].save(output_path, 'JPEG')
    return output_path

def word_to_jpg(input_path):
    doc = Document(input_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    img = Image.new('RGB', (800, 600), color=(255, 255, 255))
    output_path = os.path.join(app.config['DOWNLOAD_FOLDER'], f"{uuid.uuid4()}.jpg")
    img.save(output_path)
    return output_path

def jpg_to_word(input_path):
    img = Image.open(input_path)
    text = pytesseract.image_to_string(img)
    doc = Document()
    doc.add_paragraph(text)
    output_path = os.path.join(app.config['DOWNLOAD_FOLDER'], f"{uuid.uuid4()}.docx")
    doc.save(output_path)
    return output_path

def video_to_jpg(input_path):
    cap = cv2.VideoCapture(input_path)
    ret, frame = cap.read()
    if ret:
        output_path = os.path.join(app.config['DOWNLOAD_FOLDER'], f"{uuid.uuid4()}.jpg")
        cv2.imwrite(output_path, frame)
        cap.release()
        return output_path
    cap.release()
    return None

# Run the app
if __name__ == '__main__':
    app.run(debug=True, threaded=True)
